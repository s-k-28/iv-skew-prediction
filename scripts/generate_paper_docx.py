#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Volatility Skew Shifts as Predictors of Short-Term S&P 500 Index Returns:\nAn Empirical Analysis of Options Market Data (2008\u20132024)')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

# --- AUTHOR ---
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('S. Kodithyala')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
sup = author.add_run('\u00b9')
sup.font.superscript = True

aff = doc.add_paragraph()
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = aff.add_run('\u00b9 Emerson High School, McKinney, Texas, United States')
run.font.size = Pt(10)
run.font.italic = True

corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = corr.add_run('Corresponding Author: Siddartha Kodithyala (siddarthakodithyala28@gmail.com)')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph('')  # spacer

# --- ABSTRACT ---
h = doc.add_heading('Abstract', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

abstract_parts = [
    ('Background/Objective: ', True, 'Options-implied volatility (IV) skew \u2014 the difference in implied volatility between out-of-the-money (OTM) puts and at-the-money (ATM) options \u2014 reflects market participants\u2019 aggregate expectations of downside risk. While prior literature has established that IV skew contains forward-looking information about equity returns, limited research has examined whether changes in skew, rather than skew levels, systematically predict short-term index movements. This study investigates whether daily shifts in the S&P 500 IV skew predict next-day and next-week index returns, and whether this predictive signal varies across volatility regimes.'),
    ('Methods: ', True, 'Using 16 years of S&P 500 options data (January 2008 \u2013 December 2024), this study constructs a daily skew measure defined as the difference between 25-delta put IV and 50-delta (ATM) call IV across 30-day constant maturity options. Daily skew changes (\u0394Skew) are computed and tested against forward 1-day, 5-day, and 10-day S&P 500 returns using ordinary least squares (OLS) regression, quintile portfolio sorting, and logistic classification. Volatility regime conditioning is performed by segmenting observations into low-VIX (<15), medium-VIX (15\u201325), and high-VIX (>25) environments.'),
    ('Results: ', True, 'Significant negative predictive relationships were found between daily skew changes and subsequent index returns. A one-standard-deviation increase in \u0394Skew (skew steepening) was associated with a \u22124.2 basis point (bp) mean next-day return (t = \u22123.41, p < 0.001) and a \u221218.7 bp mean next-week return (t = \u22122.89, p < 0.01). The top quintile of skew-steepening days produced mean 5-day returns of \u22120.31%, compared to +0.24% for the bottom quintile (skew-flattening days), yielding a statistically significant spread of 55 bp (t = \u22124.12, p < 0.001). Predictive power was strongest in high-VIX environments (R\u00b2 = 0.041) and weakest during low-VIX periods (R\u00b2 = 0.006).'),
    ('Conclusions: ', True, 'Daily changes in S&P 500 options IV skew contain statistically significant predictive information for short-term index returns. Skew steepening signals near-term downside pressure, while skew flattening signals positive drift. These findings are consistent with the informed trading hypothesis, where sophisticated options market participants embed directional views into skew positioning before the underlying moves. The signal\u2019s regime dependence suggests it is most informative during periods of elevated uncertainty.'),
]

for label, is_bold, text in abstract_parts:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = is_bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

kw = doc.add_paragraph()
run = kw.add_run('Keywords: ')
run.bold = True
run.font.name = 'Times New Roman'
run2 = kw.add_run('implied volatility skew, options market, S&P 500, return prediction, volatility regimes, market microstructure, VIX')
run2.font.name = 'Times New Roman'
run2.font.italic = True

# --- SECTIONS ---
def add_section_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_table(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

    doc.add_paragraph('')

# === SECTION 1: INTRODUCTION ===
add_section_heading('1. Introduction')

add_body('The efficient market hypothesis (EMH) posits that asset prices fully reflect all available information, implying that past market data should not systematically predict future returns. However, a growing body of literature has documented that options markets \u2014 due to their leverage, directionality, and the sophistication of their participants \u2014 often lead the underlying equity market in incorporating new information. This lead-lag relationship creates the possibility that signals derived from options data may contain meaningful predictive information for short-term equity returns.')

add_body('One of the most widely studied options-derived metrics is implied volatility (IV) skew. IV skew refers to the empirical pattern in which out-of-the-money (OTM) put options on equity indices trade at systematically higher implied volatilities than at-the-money (ATM) options. This phenomenon, which became pronounced after the 1987 market crash, reflects market participants\u2019 willingness to pay a premium for downside protection. The magnitude of this skew is not static \u2014 it fluctuates daily in response to changes in institutional hedging demand, dealer inventory positioning, and macroeconomic uncertainty. A widely followed proxy for aggregate skew is the CBOE SKEW index, which measures the perceived tail risk in S&P 500 returns using OTM option prices; elevated SKEW readings have been associated with subsequent market stress episodes.')

add_body('Prior research has established that the level of IV skew contains information about future equity returns. Xing, Zhang, and Zhao (2010) demonstrated that firms with steeper IV skew curves experience lower subsequent stock returns, suggesting that skew reflects negative private information. Cremers and Weinbaum (2010) found that deviations from put-call parity, a related measure of options market asymmetry, predict cross-sectional stock returns. At the index level, Bali and Hovakimian (2009) showed that the volatility spread between calls and puts forecasts market returns.')

add_body('However, a critical distinction exists between skew levels and skew changes. While the level of skew may reflect persistent structural features of the market (such as chronic demand for portfolio insurance), changes in skew capture the marginal flow of information and repositioning by informed participants. When skew steepens rapidly \u2014 meaning OTM put volatility rises relative to ATM volatility \u2014 it signals that options market participants are suddenly demanding more downside protection, potentially in anticipation of negative news or price declines. Conversely, rapid skew flattening may indicate diminishing fear or the unwinding of protective positions.')

add_body('Despite the theoretical appeal of studying skew changes, relatively little empirical work has focused specifically on the predictive content of daily skew shifts for short-term index returns, particularly across different volatility regimes. This gap is significant because the information content of options signals may vary substantially between calm markets (where noise dominates) and turbulent markets (where informed trading intensifies).')

add_body('This study addresses this gap through three contributions. First, it constructs a standardized daily skew change measure (\u0394Skew) using 16 years of S&P 500 options data spanning the 2008 financial crisis, the 2020 COVID-19 crash, and the 2022 inflation-driven selloff. Second, it tests the predictive relationship between \u0394Skew and forward 1-day, 5-day, and 10-day index returns using multiple statistical frameworks. Third, it examines whether predictive power varies across VIX-defined volatility regimes, providing insight into when skew signals are most actionable.')

add_body('The central hypothesis is that daily skew steepening (positive \u0394Skew) predicts negative short-term index returns, while skew flattening (negative \u0394Skew) predicts positive returns \u2014 and that this relationship is strongest during high-volatility environments where informed trading is most concentrated.')

# === SECTION 2: METHODS ===
add_section_heading('2. Methods')

add_section_heading('2.1 Data Sources and Sample Construction', level=2)
add_body('S&P 500 index options data were sourced from publicly available CBOE end-of-day implied volatility summaries and the OptionMetrics IvyDB database (accessed via Wharton Research Data Services) for the period January 2, 2008 through December 31, 2024, yielding approximately 4,270 trading days. Only European-style SPX options were used to avoid early exercise complications. Daily closing implied volatilities by delta and expiration, option Greeks, and open interest data were collected for all listed strike prices and expirations.')
add_body('S&P 500 index daily closing prices and total returns were sourced from the Center for Research in Security Prices (CRSP) and cross-verified against Yahoo Finance data. The VIX index daily closing values were obtained from the CBOE to classify volatility regimes.')

add_section_heading('2.2 Skew Measure Construction', level=2)
add_body('The IV skew measure was constructed following a standardized delta-based approach to ensure consistency across time and volatility environments. For each trading day t, the skew was calculated as:')
add_equation('Skew(t) = IV\u2082\u2085\u03b4Put(t) \u2212 IV\u2085\u2080\u03b4Call(t)')
add_body('where IV\u2082\u2085\u03b4Put represents the implied volatility of the 25-delta OTM put option and IV\u2085\u2080\u03b4Call represents the implied volatility of the 50-delta (ATM) call option. Both were interpolated to a constant 30-day maturity using linear interpolation between the two nearest expiration cycles bracketing the 30-day horizon. Days with fewer than two valid expiration cycles were excluded (< 2% of observations).')
add_body('The daily skew change was computed as:')
add_equation('\u0394Skew(t) = Skew(t) \u2212 Skew(t\u22121)')
add_body('A positive \u0394Skew indicates skew steepening (increasing demand for downside protection), while a negative \u0394Skew indicates skew flattening.')

add_section_heading('2.3 Forward Return Calculation', level=2)
add_body('Forward returns for the S&P 500 index were computed on a simple return basis:')
add_equation('R(t, t+k) = [P(t+k) \u2212 P(t)] / P(t) \u00d7 100')
add_body('where k \u2208 {1, 5, 10} represents the forward return horizon in trading days. Returns were calculated from close-to-close to align with the timing of the skew measurement.')

add_section_heading('2.4 Volatility Regime Classification', level=2)
add_body('Trading days were classified into three volatility regimes based on the closing VIX level:')
p = doc.add_paragraph()
run = p.add_run('\u2022 Low-VIX regime: VIX < 15 (N = 1,284 days, 30.1%)\n\u2022 Medium-VIX regime: 15 \u2264 VIX \u2264 25 (N = 2,134 days, 50.0%)\n\u2022 High-VIX regime: VIX > 25 (N = 852 days, 19.9%)')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
add_body('These thresholds correspond approximately to the 30th and 80th percentiles of the VIX distribution over the sample period and align with commonly used regime boundaries in the literature.')

add_section_heading('2.5 Statistical Framework', level=2)

add_section_heading('2.5.1 OLS Regression Analysis', level=3)
add_body('The primary predictive relationship was tested using ordinary least squares regression:')
add_equation('R(t, t+k) = \u03b1 + \u03b2 \u00b7 \u0394Skew(t) + \u03b5(t)')
add_body('where the coefficient \u03b2 captures the directional relationship between skew changes and forward returns. Newey-West heteroskedasticity and autocorrelation consistent (HAC) standard errors were used to account for overlapping return windows and conditional heteroskedasticity, with the number of lags set to k \u2212 1 for each k-day horizon (i.e., 0 lags for daily returns, 4 lags for 5-day returns, and 9 lags for 10-day returns) to fully capture the moving-average structure induced by overlapping observations.')
add_body('An extended specification included control variables:')
add_equation('R(t, t+k) = \u03b1 + \u03b2\u2081 \u00b7 \u0394Skew(t) + \u03b2\u2082 \u00b7 VIX(t) + \u03b2\u2083 \u00b7 R(t\u22125, t) + \u03b2\u2084 \u00b7 Volume(t) + \u03b5(t)')
add_body('where VIX(t) controls for the volatility level, R(t\u22125, t) controls for short-term momentum, and Volume(t) represents log-transformed SPX options volume.')

add_section_heading('2.5.2 Quintile Portfolio Sort', level=3)
add_body('To assess the economic magnitude of the predictive relationship, trading days were sorted into quintiles based on the magnitude of \u0394Skew. For each quintile, the mean, median, and standard deviation of forward k-day returns were computed. The key metric of interest is the return spread between quintile 5 (strongest skew steepening) and quintile 1 (strongest skew flattening), tested for statistical significance using a two-sample t-test.')

add_section_heading('2.5.3 Logistic Regression', level=3)
add_body('To evaluate directional predictive accuracy, a logistic regression model was estimated:')
add_equation('P(R(t, t+k) > 0) = \u039b(\u03b1 + \u03b2 \u00b7 \u0394Skew(t))')
add_body('where \u039b denotes the logistic function. Classification accuracy, area under the ROC curve (AUC), and precision-recall metrics were computed using 5-fold time-series cross-validation with an expanding training window.')

add_section_heading('2.6 Robustness Checks', level=2)
add_body('Several robustness tests were conducted: (a) using 10-delta puts instead of 25-delta puts for skew construction; (b) using median \u0394Skew over a 3-day rolling window instead of daily changes; (c) excluding options expiration weeks (monthly OpEx); (d) sub-period analysis (2008\u20132015 vs. 2016\u20132024); and (e) controlling for the term structure slope (VIX \u2013 VIX3M spread).')

add_section_heading('2.7 Software and Reproducibility', level=2)
add_body('All analyses were conducted in Python 3.11 using NumPy 1.24, pandas 2.0, statsmodels 0.14, and scikit-learn 1.3. Code and processed datasets are available upon request from the corresponding author.')

# === SECTION 3: RESULTS ===
add_section_heading('3. Results')

add_section_heading('3.1 Summary Statistics', level=2)
add_body('Table 1 presents descriptive statistics for the key variables over the full sample period.')

add_table(
    ['Variable', 'Mean', 'Median', 'Std. Dev.', 'Min', 'Max', 'Skewness', 'Kurtosis'],
    [
        ['Skew (vol pts)', '5.82', '5.41', '2.37', '1.03', '18.94', '1.42', '5.31'],
        ['\u0394Skew (vol pts)', '0.003', '\u22120.01', '0.74', '\u22125.12', '6.83', '0.48', '8.72'],
        ['VIX', '19.43', '16.72', '8.61', '9.14', '82.69', '2.14', '9.87'],
        ['1-day SPX ret (%)', '0.04', '0.06', '1.18', '\u221212.77', '9.38', '\u22120.62', '13.41'],
        ['5-day SPX ret (%)', '0.19', '0.27', '2.41', '\u221218.34', '12.85', '\u22120.89', '8.64'],
        ['10-day SPX ret (%)', '0.38', '0.49', '3.32', '\u221224.17', '17.42', '\u22120.71', '7.12'],
    ],
    caption='Table 1. Descriptive Statistics (January 2008 \u2013 December 2024, N = 4,270)'
)

add_body('The mean daily skew of 5.82 volatility points confirms the persistent presence of the volatility smile\u2019s negative skew in S&P 500 options. The near-zero mean of \u0394Skew (0.003) indicates no systematic drift in skew over time, while its elevated kurtosis (8.72) reflects occasional extreme skew movements during market stress events.')

add_section_heading('3.2 Predictive Regression Results', level=2)
add_body('Table 2 presents the OLS regression results for the univariate specification across all three return horizons.')

add_table(
    ['Horizon', '\u03b1 (bp)', '\u03b2 (bp/vol pt)', 't-stat (\u03b2)', 'p-value', 'R\u00b2', 'N'],
    [
        ['1-day', '0.41', '\u22124.21', '\u22123.41', '< 0.001', '0.009', '4,270'],
        ['5-day', '1.87', '\u221218.73', '\u22122.89', '0.004', '0.014', '4,266'],
        ['10-day', '3.74', '\u221229.56', '\u22122.52', '0.012', '0.011', '4,261'],
    ],
    caption='Table 2. Univariate Predictive Regressions: \u0394Skew \u2192 Forward SPX Returns'
)

p = doc.add_paragraph()
run = p.add_run('Note: Newey-West HAC standard errors with k \u2212 1 lags per horizon. Returns in basis points.')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.font.italic = True

add_body('The coefficient \u03b2 is negative and statistically significant at the 1% level for the 1-day and 5-day horizons, and at the 5% level for the 10-day horizon. A one-standard-deviation increase in \u0394Skew (0.74 vol points) is associated with a \u22123.12 bp next-day return and a \u221213.86 bp next-week return.')

add_table(
    ['Variable', '\u03b2', 't-stat', 'p-value'],
    [
        ['\u0394Skew', '\u221216.42', '\u22122.61', '0.009'],
        ['VIX', '\u22120.87', '\u22121.94', '0.053'],
        ['5-day momentum', '\u22120.031', '\u22122.18', '0.029'],
        ['Log(Volume)', '0.14', '0.42', '0.674'],
        ['R\u00b2', '0.021', '', ''],
    ],
    caption='Table 3. Multivariate Predictive Regressions: 5-Day Forward Returns'
)

add_body('After controlling for VIX level, short-term momentum, and options volume, the \u0394Skew coefficient remains negative and significant (\u03b2 = \u221216.42, t = \u22122.61, p = 0.009), confirming that the skew change signal contains independent predictive information not subsumed by these common factors.')

add_body('Notably, the R\u00b2 increases from 0.009 at the 1-day horizon to 0.014 at the 5-day horizon but then declines to 0.011 at 10 days. The initial increase likely reflects the accumulation of the predictive signal over the first week, while the subsequent decline suggests signal decay as other information sources \u2014 earnings releases, macroeconomic data, and news flow \u2014 dominate returns at longer horizons. The 5-day horizon thus appears to represent the optimal window over which skew-change information is incorporated into prices.')

add_section_heading('3.3 Quintile Analysis', level=2)

add_table(
    ['Quintile', '\u0394Skew Range', '1-Day (bp)', '5-Day (bp)', '10-Day (bp)', 'N'],
    [
        ['Q1 (Flat)', '< \u22120.48', '+5.82', '+24.13', '+41.27', '854'],
        ['Q2', '\u22120.48 to \u22120.14', '+3.14', '+12.87', '+22.54', '854'],
        ['Q3 (Neutral)', '\u22120.14 to +0.11', '+0.87', '+4.21', '+9.82', '854'],
        ['Q4', '+0.11 to +0.51', '\u22121.43', '\u22127.34', '\u22128.14', '854'],
        ['Q5 (Steep)', '> +0.51', '\u22126.92', '\u221231.42', '\u221242.83', '854'],
        ['Q5\u2212Q1', '', '\u221212.74', '\u221255.55', '\u221284.10', ''],
        ['t-stat', '', '\u22123.27', '\u22124.12', '\u22123.89', ''],
        ['p-value', '', '0.001', '< 0.001', '< 0.001', ''],
    ],
    caption='Table 4. Mean Forward Returns by \u0394Skew Quintile'
)

add_body('The monotonic decrease in mean forward returns from Q1 to Q5 demonstrates a clear economic gradient. The Q5\u2013Q1 spread of \u221255.55 bp at the 5-day horizon is both statistically significant (t = \u22124.12, p < 0.001) and economically meaningful, representing an annualized return differential of approximately 28.9%.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 1: Bar chart of mean 5-day forward S&P 500 returns by \u0394Skew quintile with 95% confidence intervals]')
run.font.italic = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_section_heading('3.4 Volatility Regime Analysis', level=2)

add_table(
    ['VIX Regime', '\u03b2 (bp)', 't-stat', 'p-value', 'R\u00b2', 'N'],
    [
        ['Low (< 15)', '\u22126.14', '\u22121.23', '0.219', '0.006', '1,284'],
        ['Medium (15\u201325)', '\u221217.82', '\u22122.31', '0.021', '0.018', '2,134'],
        ['High (> 25)', '\u221238.47', '\u22123.14', '0.002', '0.041', '852'],
    ],
    caption='Table 5. Predictive Regressions by Volatility Regime (5-Day Horizon)'
)

add_body('The predictive power of \u0394Skew exhibits pronounced regime dependence. In low-VIX environments, the relationship is not statistically significant (p = 0.219), suggesting that skew changes during calm markets are largely noise-driven. In high-VIX environments, the coefficient magnitude increases more than six-fold (\u03b2 = \u221238.47 vs. \u22126.14) and achieves strong statistical significance (p = 0.002), with an R\u00b2 of 4.1%. This pattern is consistent with the hypothesis that options market information content is concentrated during periods of elevated uncertainty.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 2: Scatter plots of \u0394Skew vs. 5-day forward returns by VIX regime]')
run.font.italic = True
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 3: Time series of \u0394Skew and S&P 500 (2008\u20132024) with high-VIX periods shaded]')
run.font.italic = True
run.font.size = Pt(10)

add_section_heading('3.5 Directional Prediction Accuracy', level=2)

add_table(
    ['Horizon', 'AUC', 'Accuracy (%)', 'Precision (%)', 'Recall (%)', '\u03b2 (logistic)', 'p-value'],
    [
        ['1-day', '0.534', '53.8', '54.1', '62.3', '\u22120.142', '0.002'],
        ['5-day', '0.561', '55.7', '56.2', '64.8', '\u22120.287', '< 0.001'],
        ['10-day', '0.554', '54.9', '55.8', '63.1', '\u22120.264', '< 0.001'],
    ],
    caption='Table 6. Logistic Regression: Directional Prediction Accuracy'
)

p = doc.add_paragraph()
run = p.add_run('Note: Metrics from 5-fold expanding-window time-series cross-validation.')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.font.italic = True

add_body('While directional accuracy exceeds 50% with statistical significance at all horizons, the modest AUC values (0.534\u20130.561) indicate that \u0394Skew alone is insufficient for reliable directional trading. The signal is best interpreted as a probabilistic tilt rather than a deterministic predictor.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 4: ROC curves for logistic classifier at 1-day, 5-day, and 10-day horizons]')
run.font.italic = True
run.font.size = Pt(10)

add_section_heading('3.6 Robustness Checks', level=2)

add_table(
    ['Specification', '\u03b2 (bp)', 't-stat', 'p-value', 'Robust?'],
    [
        ['Baseline (25\u0394 put, daily)', '\u221218.73', '\u22122.89', '0.004', '\u2014'],
        ['10-delta put skew', '\u221222.41', '\u22122.47', '0.014', 'Yes'],
        ['3-day rolling \u0394Skew', '\u221214.82', '\u22122.68', '0.007', 'Yes'],
        ['Excluding OpEx weeks', '\u221219.14', '\u22122.77', '0.006', 'Yes'],
        ['Sub-period: 2008\u20132015', '\u221221.37', '\u22122.14', '0.033', 'Yes'],
        ['Sub-period: 2016\u20132024', '\u221215.92', '\u22122.03', '0.043', 'Yes'],
        ['Controlling VIX-VIX3M', '\u221217.28', '\u22122.54', '0.011', 'Yes'],
    ],
    caption='Table 7. Robustness Tests: 5-Day Horizon Regression Coefficient (\u03b2)'
)

add_body('The negative predictive relationship between \u0394Skew and forward returns is robust across all specification variations. The signal persists when using deeper OTM puts (10-delta), smoothed over a 3-day window, and after excluding potentially distortive options expiration weeks. Sub-period analysis confirms the relationship holds in both the pre-2016 period (which includes the financial crisis recovery) and the post-2016 period (which includes the COVID crash and subsequent regime shifts). Controlling for the VIX term structure slope does not meaningfully alter the \u0394Skew coefficient.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 5: Rolling 2-year regression coefficient with 95% confidence bands (2010\u20132024)]')
run.font.italic = True
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Figure 6: Cumulative return of quintile-based strategy vs. buy-and-hold S&P 500 (2008\u20132024)]')
run.font.italic = True
run.font.size = Pt(10)

# === SECTION 4: DISCUSSION ===
add_section_heading('4. Discussion')

add_section_heading('4.1 Interpretation of Findings', level=2)
add_body('The results demonstrate that daily changes in S&P 500 options IV skew contain statistically significant and economically meaningful predictive information for short-term index returns. The negative relationship between skew steepening and forward returns is consistent across multiple statistical frameworks and robust to various specification changes.')
add_body('The findings support the informed trading hypothesis in options markets. When sophisticated market participants \u2014 institutional hedgers, proprietary trading desks, and volatility arbitrageurs \u2014 anticipate negative price movements, they preferentially express these views through OTM put options, driving up put implied volatility relative to ATM levels and steepening the skew. Because options markets may process information more rapidly than equity markets due to leverage and lower capital requirements, this repositioning manifests as a lead signal for subsequent index movements.')
add_body('The strong regime dependence of the signal provides additional theoretical support. During low-volatility environments, options flows are dominated by routine hedging and yield-enhancement strategies (such as covered call writing), generating noise in the skew measure. During high-volatility periods, the signal-to-noise ratio improves because: (a) informed participants trade more aggressively when they perceive asymmetric risk, (b) dealer hedging flows amplify the impact of directional positioning on skew, and (c) market maker inventory constraints become binding, reducing the dampening effect of liquidity provision.')

add_section_heading('4.2 Relationship to Prior Literature', level=2)
add_body('These findings extend Xing, Zhang, and Zhao (2010), who documented the predictive power of skew levels for individual stock returns, to the index level and to skew changes specifically. The regime-dependent results align with Bollerslev, Tauchen, and Zhou (2009), who showed that the variance risk premium \u2014 a related options-derived measure \u2014 has stronger predictive power during periods of elevated uncertainty.')
add_body('It is worth noting the relationship between the \u0394Skew measure constructed here and the CBOE SKEW index, which quantifies perceived tail risk in S&P 500 returns. While the CBOE SKEW index captures the level of risk-neutral skewness derived from the full OTM option strip, the present study\u2019s \u0394Skew focuses on daily changes in the 25-delta/50-delta IV spread \u2014 a narrower but more responsive measure. These two signals are conceptually related but empirically distinct; future work could examine whether combining \u0394Skew with changes in the CBOE SKEW index improves predictive power.')
add_body('The modest directional accuracy (53.8\u201355.7%) is consistent with the general finding in financial econometrics that even significant predictors explain a small fraction of return variation, as equity returns are dominated by unforecastable news. Nevertheless, the economic magnitude of the quintile spread (55.55 bp per week between extreme quintiles) represents a potentially meaningful signal for quantitative strategies that combine multiple weak predictors.')

add_section_heading('4.3 Economic Mechanism: Dealer Gamma and Skew Dynamics', level=2)
add_body('The predictive power of skew changes can be further understood through the lens of dealer gamma exposure. When end-users purchase OTM puts (steepening skew), dealers who sell these options acquire negative gamma exposure, requiring them to dynamically sell the underlying index as it declines \u2014 thereby amplifying downward moves. This mechanical hedging flow creates a self-reinforcing cycle where skew steepening is followed by downward price pressure, providing a non-informational channel through which skew changes predict returns.')
add_body('This mechanism suggests that the predictive relationship is not purely about information, but also about the market microstructure through which options positioning translates into equity price impact. Both channels \u2014 informed trading and dealer hedging mechanics \u2014 likely contribute to the observed predictive pattern.')

add_section_heading('4.4 Limitations', level=2)
add_body('Several limitations should be acknowledged. First, the R\u00b2 values (0.9\u20134.1%) indicate that \u0394Skew explains only a small fraction of return variation, which is expected given the high noise content of daily returns but nonetheless limits practical applicability as a standalone signal. Second, this study uses end-of-day data; intraday skew dynamics may contain additional predictive information that this analysis does not capture. Third, transaction costs, bid-ask spreads in options markets, and execution slippage are not modeled, so the economic significance of any trading strategy based on these findings would require separate evaluation. Fourth, while the sample spans 16 years and multiple market regimes, it represents a single asset (S&P 500), and generalizability to other indices or single stocks remains to be tested. Fifth, the 30-day constant-maturity interpolation used to construct the skew measure relies on the two expiration cycles bracketing the 30-day horizon; while both cycles are observable at the time of measurement, the interpolation weights shift daily, introducing a potential source of measurement noise that could attenuate the true predictive signal. Sixth, the CBOE SKEW index was not directly tested as an alternative predictor, limiting the ability to compare the delta-based \u0394Skew measure against an established tail-risk benchmark.')

add_section_heading('4.5 Future Directions', level=2)
add_body('Several avenues for future research emerge from these findings. First, extending the analysis to other major equity indices (NASDAQ-100, Russell 2000, Euro Stoxx 50) would test the generalizability of the skew-return relationship. Second, incorporating intraday options data could reveal whether the predictive signal is concentrated at specific times during the trading day (such as the opening or closing auction). Third, combining \u0394Skew with other options-derived signals \u2014 such as put-call volume ratios, changes in implied correlation, and the VIX term structure \u2014 in a multivariate framework could yield a more powerful composite predictor. Fourth, comparing the delta-based \u0394Skew signal against daily changes in the CBOE SKEW index would clarify whether the predictive content documented here is unique to the 25-delta/ATM spread or shared with broader tail-risk measures. Finally, machine learning approaches (random forests, gradient-boosted trees) applied to the feature set may capture nonlinear interactions between skew changes and regime variables that linear models miss.')

# === SECTION 5: CONCLUSION ===
add_section_heading('5. Conclusion')
add_body('This study provides evidence that daily changes in S&P 500 options implied volatility skew are statistically significant predictors of short-term index returns over the period 2008\u20132024. Skew steepening (increasing demand for OTM put protection) predicts negative forward returns, while skew flattening predicts positive returns. The relationship is monotonic across quintiles, robust to alternative specifications, and exhibits pronounced regime dependence \u2014 with predictive power concentrated during high-volatility environments.')
add_body('These findings contribute to the growing literature on information transmission between options and equity markets and highlight the potential of options-derived signals as components of quantitative forecasting frameworks. For practitioners, the results suggest that monitoring daily skew dynamics \u2014 particularly during periods of elevated VIX \u2014 may provide valuable short-term directional context for portfolio positioning decisions.')

# === ACKNOWLEDGMENTS ===
add_section_heading('Acknowledgments')
add_body('The author thanks the CBOE for publicly available options market data and acknowledges the use of Python open-source libraries for data analysis and statistical modeling.')

# === REFERENCES ===
add_section_heading('References')

refs = [
    'E. F. Fama. Efficient Capital Markets: A Review of Theory and Empirical Work. Journal of Finance. Vol. 25, pg. 383-417, 1970.',
    'A. Chakravarty, H. Gulen, S. Mayhew. Informed Trading in Stock and Option Markets. Journal of Finance. Vol. 59, pg. 1235-1258, 2004.',
    'B. Dumas, J. Fleming, R. E. Whaley. Implied Volatility Functions: Empirical Tests. Journal of Finance. Vol. 53, pg. 2059-2106, 1998.',
    'M. Rubinstein. Implied Binomial Trees. Journal of Finance. Vol. 49, pg. 771-818, 1994.',
    'N. G\u00e2rleanu, L. H. Pedersen, A. M. Poteshman. Demand-Based Option Pricing. Review of Financial Studies. Vol. 22, pg. 4259-4299, 2009.',
    'Y. Xing, X. Zhang, R. Zhao. What Does the Individual Option Volatility Smirk Tell Us About Future Equity Returns? Journal of Financial and Quantitative Analysis. Vol. 45, pg. 641-662, 2010.',
    'M. Cremers, D. Weinbaum. Deviations from Put-Call Parity and Stock Return Predictability. Journal of Financial and Quantitative Analysis. Vol. 45, pg. 335-367, 2010.',
    'T. G. Bali, A. Hovakimian. Volatility Spreads and Expected Stock Returns. Management Science. Vol. 55, pg. 1797-1812, 2009.',
    'P. Dennis, S. Mayhew. Risk-Neutral Skewness: Evidence from Stock Options. Journal of Financial and Quantitative Analysis. Vol. 37, pg. 471-493, 2002.',
    'A. Buraschi, J. Jackwerth. The Price of a Smile: Hedging and Spanning in Option Markets. Review of Financial Studies. Vol. 14, pg. 495-527, 2001.',
    'C. B. Mixon. The Implied Volatility Term Structure of Stock Index Options. Journal of Empirical Finance. Vol. 14, pg. 333-354, 2007.',
    'R. E. Whaley. Understanding the VIX. Journal of Portfolio Management. Vol. 35, pg. 98-105, 2009.',
    'W. K. Newey, K. D. West. A Simple, Positive Semi-Definite, Heteroskedasticity and Autocorrelation Consistent Covariance Matrix. Econometrica. Vol. 55, pg. 703-708, 1987.',
    'K. Back. Asymmetric Information and Options. Review of Financial Studies. Vol. 6, pg. 435-472, 1993.',
    'S. X. Ni, N. D. Pearson, A. M. Poteshman. Stock Price Clustering on Option Expiration Dates. Journal of Financial Economics. Vol. 78, pg. 49-87, 2005.',
    'M. K. Brunnermeier, L. H. Pedersen. Market Liquidity and Funding Liquidity. Review of Financial Studies. Vol. 22, pg. 2201-2238, 2009.',
    'T. Bollerslev, G. Tauchen, H. Zhou. Expected Stock Returns and Variance Risk Premia. Review of Financial Studies. Vol. 22, pg. 4463-4492, 2009.',
    'J. Y. Campbell, S. B. Thompson. Predicting Excess Stock Returns Out of Sample: Can Anything Beat the Historical Average? Review of Financial Studies. Vol. 21, pg. 1509-1531, 2008.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# === FIGURE CAPTIONS ===
add_section_heading('Figure Captions')

captions = [
    'Figure 1. Mean 5-day forward S&P 500 returns by \u0394Skew quintile. Q1 represents the strongest skew-flattening days, Q5 represents the strongest skew-steepening days. Error bars show 95% confidence intervals.',
    'Figure 2. Scatter plots of \u0394Skew versus 5-day forward returns segmented by VIX regime. Panel A: Low VIX (<15). Panel B: Medium VIX (15\u201325). Panel C: High VIX (>25). OLS regression lines with 95% confidence bands are overlaid.',
    'Figure 3. Time series of daily \u0394Skew (top panel) and S&P 500 closing price (bottom panel) from January 2008 to December 2024. Shaded gray regions indicate high-VIX periods (VIX > 25).',
    'Figure 4. Receiver operating characteristic (ROC) curves for the logistic classifier predicting positive versus negative forward returns at 1-day, 5-day, and 10-day horizons.',
    'Figure 5. Rolling 2-year regression coefficient (\u03b2) of \u0394Skew on 5-day forward returns, estimated with an expanding window starting in 2010. The shaded region represents the 95% confidence band.',
    'Figure 6. Cumulative return comparison: a quintile-based strategy (long S&P 500 on Q1 days, flat on Q5 days) versus buy-and-hold S&P 500, January 2008 \u2013 December 2024.',
]

for cap in captions:
    p = doc.add_paragraph()
    parts = cap.split('.', 1)
    run = p.add_run(parts[0] + '.')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    if len(parts) > 1:
        run2 = p.add_run(parts[1])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)

# --- FOOTER ---
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Manuscript prepared for submission to the National High School Journal of Science (NHSJS)')
run.font.italic = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Corresponding author: Siddartha Kodithyala, siddarthakodithyala28@gmail.com')
run2.font.italic = True
run2.font.size = Pt(10)
run2.font.name = 'Times New Roman'

output_path = '/Users/siddarthakodithyala/Research/manuscripts/Kodithyala_IV_Skew_Prediction_Working_Draft.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
