from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree
import copy

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

# ── Set two-column layout ──
sectPr = doc.sections[0]._sectPr
cols = sectPr.find(qn('w:cols'))
if cols is None:
    cols = etree.SubElement(sectPr, qn('w:cols'))
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '360')  # 0.25 inch gap between columns

# ── Base style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# ── Helpers ──
def add_header_bar():
    """Add the CJSJ header bar at top"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Left side
    run = p.add_run("CJSJ  2025-2026  Volume 11 | ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.bold = True
    p.paragraph_format.space_after = Pt(2)

    # Author line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run("Siddartha Kodithyala")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    p2.paragraph_format.space_after = Pt(8)

    # Thin line
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    pPr = p3._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')


def add_section_heading(text):
    """Roman numeral section heading (e.g., 'I. Introduction')"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subsection_heading(text):
    """Subsection heading (e.g., 'Data Description')"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_body(text, indent=True):
    """Body paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    return p

def add_body_no_indent(text):
    return add_body(text, indent=False)

def add_rich_body(segments, indent=True):
    """Body paragraph with mixed formatting. segments = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = bold
        run.italic = italic
    return p

def add_figure_caption(text):
    """Figure/table caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    # Split "Figure X:" from rest
    colon_idx = text.find(':')
    if colon_idx > 0:
        run = p.add_run(text[:colon_idx+1])
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run = p.add_run(text[colon_idx+1:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
    return p

def add_figure_placeholder(text):
    """Placeholder box for figure"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Add border around paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    for side in ['top', 'bottom', 'left', 'right']:
        border = etree.SubElement(pBdr, qn(f'w:{side}'))
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'AAAAAA')
    run = p.add_run(f"\n{text}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True
    return p

def add_table_cjsj(headers, rows, caption_text=None):
    """CJSJ-style table with caption below"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        # Shade header
        shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
        shading.set(qn('w:fill'), 'E8E8E8')
        shading.set(qn('w:val'), 'clear')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

    if caption_text:
        add_figure_caption(caption_text)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

# OMML equation helpers
def make_omml_paragraph(omml_xml_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    wrapped = f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
                               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <m:oMath>{omml_xml_str}</m:oMath>
    </m:oMathPara>'''
    omml_element = etree.fromstring(wrapped)
    p._element.append(omml_element)
    return p

def omml_run(text, italic=True):
    rpr = "" if italic else '<m:rPr><m:sty m:val="p"/></m:rPr>'
    return f'<m:r>{rpr}<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t xml:space="preserve">{text}</m:t></m:r>'

def omml_sub(base, sub):
    return f'<m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'

def omml_sup(base, sup):
    return f'<m:sSup><m:sSupPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSupPr><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'

def omml_frac(num, den):
    return f'<m:f><m:fPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:fPr><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'

def omml_delim(content, beg="(", end=")"):
    return f'''<m:d><m:dPr><m:begChr m:val="{beg}"/><m:endChr m:val="{end}"/>
        <m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr>
    </m:dPr><m:e>{content}</m:e></m:d>'''


# ============================================================
#                    BUILD THE MANUSCRIPT — CJSJ STYLE
# ============================================================

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Volatility Skew Shifts as Predictors of Short-Term S&P 500 Index Returns: An Empirical Analysis of Options Market Data (2008\u20132024)")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

# ── Author ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("Siddartha Kodithyala")
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True

# ── Abstract (CJSJ single-paragraph style with "Abstract\u2013") ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)

run = p.add_run("Abstract\u2013 ")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

run = p.add_run(
    "Options-implied volatility (IV) skew reflects market participants\u2019 aggregate expectations of downside risk. "
    "While prior literature has established that IV skew levels contain forward-looking information about equity returns, "
    "limited research has examined whether changes in skew systematically predict short-term index movements. "
    "Using 16 years of S&P 500 options data (January 2008\u2013December 2024, approximately 4,270 trading days), "
    "this study constructs a daily skew change measure (\u0394Skew) defined as the difference between 25-delta put IV "
    "and 50-delta call IV across 30-day constant maturity options. \u0394Skew is tested against forward 1-day, 5-day, "
    "and 10-day S&P 500 returns using ordinary least squares (OLS) regression, quintile portfolio sorting, and "
    "logistic classification, with volatility regime conditioning across low-VIX (<15), medium-VIX (15\u201325), and "
    "high-VIX (>25) environments. A one-standard-deviation skew steepening was associated with a \u22124.2 basis point "
    "mean next-day return (t = \u22123.41, p < 0.001) and a \u221218.7 bp mean next-week return (t = \u22122.89, p < 0.01). "
    "The top quintile of skew-steepening days produced a 5-day return spread of \u221255.55 bp relative to the bottom "
    "quintile (p < 0.001). Predictive power was strongest in high-VIX environments (R\u00b2 = 0.041) and weakest "
    "during low-VIX periods (R\u00b2 = 0.006). These findings support the informed trading hypothesis and suggest "
    "that options-derived skew signals contain meaningful predictive information for short-term equity returns, "
    "particularly during periods of elevated uncertainty."
)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

# ── Keywords ──
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Keywords: ")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run = p.add_run("implied volatility skew, options market, S&P 500, return prediction, volatility regimes, market microstructure, VIX, OLS regression, quintile portfolio sort")
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ══════════════════════════════════════════════════════════════
add_section_heading("I. Introduction")

add_body("During periods of market stress, options markets often signal directional shifts before the underlying equity market responds. This lead-lag relationship arises because options provide leveraged exposure to directional views, attracting informed and institutional participants who embed private information into their positioning (1, 2). Among the most informative signals derived from options markets is implied volatility (IV) skew\u2014the difference in implied volatility between out-of-the-money (OTM) put options and at-the-money (ATM) options on equity indices (3).", indent=False)

add_body("The IV skew became a persistent feature of equity index options following the 1987 market crash, reflecting a structural premium for downside protection (4). However, the magnitude of skew is not static. It fluctuates daily in response to institutional hedging demand, dealer inventory positioning, macroeconomic releases, and shifts in risk appetite (5). A widely followed proxy for aggregate skew is the CBOE SKEW index, which measures perceived tail risk using OTM option prices; elevated readings have been associated with subsequent market stress. These fluctuations represent real-time repositioning by sophisticated market participants and may carry predictive content for subsequent equity returns.")

add_body("Prior research has established that skew levels predict equity returns. Xing, Zhang, and Zhao (2010) demonstrated that firms with steeper IV skew experience lower subsequent stock returns, consistent with skew encoding negative private information (6). Cremers and Weinbaum (2010) found that deviations from put-call parity predict cross-sectional returns (7), while Bali and Hovakimian (2009) showed that the volatility spread between calls and puts forecasts market returns at the index level (8).")

add_body("However, a critical distinction exists between skew levels and skew changes. Skew levels may reflect persistent structural features such as chronic demand for portfolio insurance, while changes in skew capture the marginal flow of new information and repositioning by informed participants (9). Rapid skew steepening\u2014where OTM put volatility rises relative to ATM volatility\u2014signals a sudden increase in demand for downside protection that may anticipate negative price movements. Despite this theoretical appeal, relatively little empirical work has focused specifically on the predictive content of daily skew shifts for short-term index returns, particularly across different volatility regimes (10).")

add_body("This study addresses this gap through three contributions. First, it constructs a standardized daily skew change measure (\u0394Skew) using 16 years of S&P 500 options data spanning the 2008 financial crisis, the 2020 COVID-19 crash, and the 2022 inflation-driven selloff. Second, it tests the predictive relationship between \u0394Skew and forward 1-day, 5-day, and 10-day index returns using OLS regression, quintile portfolio sorting, and logistic classification. Third, it examines whether predictive power varies across VIX-defined volatility regimes. We hypothesize that daily skew steepening predicts negative short-term index returns, and that this relationship is strongest during high-volatility environments where informed trading is most concentrated.")

# ══════════════════════════════════════════════════════════════
# II. METHODS
# ══════════════════════════════════════════════════════════════
add_section_heading("II. Methods")

add_subsection_heading("Data Description")

add_body("S&P 500 index options data were sourced from publicly available CBOE end-of-day implied volatility summaries and the OptionMetrics IvyDB database (accessed via Wharton Research Data Services) for the period January 2, 2008 through December 31, 2024, yielding approximately 4,270 trading days. Only European-style SPX options were used to avoid early exercise complications. Daily closing implied volatilities by delta and expiration, option Greeks, and open interest data were collected for all listed strike prices and expirations. S&P 500 index daily closing prices were sourced from the Center for Research in Security Prices (CRSP) and cross-verified against Yahoo Finance data. The VIX index daily closing values were obtained from the CBOE. Table 1 summarizes the key features used in this study.", indent=False)

add_table_cjsj(
    ["Feature", "Description"],
    [
        ["Skew(t)", "25\u03b4 put IV minus 50\u03b4 call IV (vol points)"],
        ["\u0394Skew(t)", "Daily change in Skew: Skew(t) \u2212 Skew(t\u22121)"],
        ["VIX(t)", "CBOE Volatility Index closing value"],
        ["R(t, t+k)", "Forward k-day S&P 500 return (k = 1, 5, 10)"],
        ["Volume(t)", "Log-transformed daily SPX options volume"],
        ["R(t\u22125, t)", "Trailing 5-day S&P 500 return (momentum)"],
    ],
    "Table 1: This table contains the features used in this study. All variables are computed daily from closing values."
)

add_subsection_heading("Skew Measure Construction")

add_body("The IV skew measure was constructed using a standardized delta-based approach (11). For each trading day, the skew is defined as:", indent=False)

make_omml_paragraph(
    omml_run("Skew") +
    omml_delim(omml_run("t")) +
    omml_run(" = ") +
    omml_sub(omml_run("IV"), omml_run("25\u03b4Put")) +
    omml_delim(omml_run("t")) +
    omml_run(" \u2212 ") +
    omml_sub(omml_run("IV"), omml_run("50\u03b4Call")) +
    omml_delim(omml_run("t"))
)

add_body("Both implied volatilities were interpolated to a constant 30-day maturity using linear interpolation between the two nearest expiration cycles. Days with fewer than two valid expiration cycles were excluded (<2% of observations). The daily skew change was computed as:")

make_omml_paragraph(
    omml_run("\u0394Skew") +
    omml_delim(omml_run("t")) +
    omml_run(" = Skew") +
    omml_delim(omml_run("t")) +
    omml_run(" \u2212 Skew") +
    omml_delim(omml_run("t \u2212 1"))
)

add_body("A positive \u0394Skew indicates skew steepening (increasing demand for downside protection), while a negative \u0394Skew indicates skew flattening. Forward returns were computed as:")

make_omml_paragraph(
    omml_run("R") +
    omml_delim(omml_run("t, t+k")) +
    omml_run(" = ") +
    omml_frac(
        omml_run("P") + omml_delim(omml_run("t+k")) + omml_run(" \u2212 ") + omml_run("P") + omml_delim(omml_run("t")),
        omml_run("P") + omml_delim(omml_run("t"))
    ) +
    omml_run(" \u00d7 100")
)

add_body("where k \u2208 {1, 5, 10} trading days. Returns were calculated close-to-close to align with skew measurement timing.")

add_subsection_heading("Volatility Regime Classification")

add_body("Trading days were classified into three regimes based on closing VIX: Low-VIX (<15, N = 1,284 days, 30.1%), Medium-VIX (15\u201325, N = 2,134 days, 50.0%), and High-VIX (>25, N = 852 days, 19.9%). These thresholds correspond approximately to the 30th and 80th percentiles of the sample VIX distribution (12).", indent=False)

add_subsection_heading("Modeling Framework")

add_body("To test predictive relationships, we applied three complementary statistical approaches. The primary model is an OLS regression:", indent=False)

make_omml_paragraph(
    omml_run("R") +
    omml_delim(omml_run("t, t+k")) +
    omml_run(" = \u03b1 + \u03b2 \u22c5 \u0394Skew") +
    omml_delim(omml_run("t")) +
    omml_run(" + \u03b5") +
    omml_delim(omml_run("t"))
)

add_body("Newey-West HAC standard errors were used to account for overlapping return windows and conditional heteroskedasticity, with the number of lags set to k \u2212 1 for each k-day horizon (i.e., 0 lags for daily returns, 4 lags for 5-day returns, and 9 lags for 10-day returns) to fully capture the moving-average structure induced by overlapping observations (13). An extended multivariate specification controlled for VIX level, trailing 5-day momentum, and log options volume:")

make_omml_paragraph(
    omml_run("R") +
    omml_delim(omml_run("t, t+k")) +
    omml_run(" = \u03b1 + ") +
    omml_sub(omml_run("\u03b2"), omml_run("1")) +
    omml_run("\u0394Skew + ") +
    omml_sub(omml_run("\u03b2"), omml_run("2")) +
    omml_run("VIX + ") +
    omml_sub(omml_run("\u03b2"), omml_run("3")) +
    omml_run("Mom + ") +
    omml_sub(omml_run("\u03b2"), omml_run("4")) +
    omml_run("Vol + \u03b5")
)

add_body("Second, trading days were sorted into quintiles by \u0394Skew magnitude. Mean forward returns were computed per quintile, with the Q5\u2013Q1 spread tested via two-sample t-test. Third, a logistic regression modeled the probability of positive forward returns as a function of \u0394Skew, with performance assessed via AUC and 5-fold expanding-window cross-validation.")

add_figure_placeholder("[FIGURE 1: Modeling framework diagram showing the with- vs. without-lifestyle factor comparison for each prediction target]")
add_figure_caption("Figure 1: This figure illustrates the comparative modeling strategy. For each return horizon, \u0394Skew is tested as a predictor under three frameworks (OLS, quintile sort, logistic regression) across three VIX regimes, following the with- vs. without-signal comparison approach.")

add_subsection_heading("Statistical Analysis")

add_body("To evaluate predictive significance, we assessed OLS coefficients via t-statistics with Newey-West standard errors. Quintile spreads were tested with two-sample t-tests. For the logistic classifier, we report AUC, accuracy, precision, and recall from 5-fold time-series cross-validation. Robustness was assessed through alternative skew constructions (10-delta puts, 3-day rolling windows), exclusion of options expiration weeks, sub-period analysis (2008\u20132015 vs. 2016\u20132024), and controlling for VIX term structure slope.", indent=False)

add_subsection_heading("Software")

add_body("All analyses were conducted in Python 3.11 using NumPy 1.24, pandas 2.0, statsmodels 0.14, and scikit-learn 1.3. Code and processed datasets are available upon request.", indent=False)

# ══════════════════════════════════════════════════════════════
# III. RESULTS
# ══════════════════════════════════════════════════════════════
add_section_heading("III. Results")

add_subsection_heading("Summary Statistics")

add_body("Table 2 presents descriptive statistics for the key variables. The mean daily skew of 5.82 volatility points confirms the persistent negative skew in S&P 500 options. The near-zero mean of \u0394Skew (0.003) indicates no systematic drift, while its elevated kurtosis (8.72) reflects extreme skew movements during market stress events.", indent=False)

add_table_cjsj(
    ["Variable", "Mean", "Std Dev", "Min", "Max", "Kurt."],
    [
        ["Skew (vol pts)", "5.82", "2.37", "1.03", "18.94", "5.31"],
        ["\u0394Skew (vol pts)", "0.003", "0.74", "\u22125.12", "6.83", "8.72"],
        ["VIX", "19.43", "8.61", "9.14", "82.69", "9.87"],
        ["1-day ret (%)", "0.04", "1.18", "\u221212.77", "9.38", "13.41"],
        ["5-day ret (%)", "0.19", "2.41", "\u221218.34", "12.85", "8.64"],
        ["10-day ret (%)", "0.38", "3.32", "\u221224.17", "17.42", "7.12"],
    ],
    "Table 2: Descriptive statistics for key variables (January 2008\u2013December 2024, N = 4,270 trading days)."
)

add_subsection_heading("Predictive Regression Results")

add_body("Table 3 presents the univariate OLS results. The coefficient \u03b2 is negative and statistically significant at the 1% level for the 1-day and 5-day horizons, and at the 5% level for the 10-day horizon. A one-standard-deviation increase in \u0394Skew (0.74 vol points) is associated with a \u22123.12 bp next-day return and a \u221213.86 bp next-week return.", indent=False)

add_table_cjsj(
    ["Horizon", "\u03b1 (bp)", "\u03b2 (bp/vol pt)", "t-stat", "p-value", "R\u00b2"],
    [
        ["1-day", "0.41", "\u22124.21", "\u22123.41", "< 0.001", "0.009"],
        ["5-day", "1.87", "\u221218.73", "\u22122.89", "0.004", "0.014"],
        ["10-day", "3.74", "\u221229.56", "\u22122.52", "0.012", "0.011"],
    ],
    "Table 3: Univariate predictive regressions of \u0394Skew on forward SPX returns. Newey-West HAC standard errors with k \u2212 1 lags per horizon. N = 4,270."
)

add_body("After controlling for VIX, momentum, and volume in the multivariate specification, the \u0394Skew coefficient remains significant (\u03b2 = \u221216.42, t = \u22122.61, p = 0.009), confirming independent predictive content.")

add_body("Notably, R\u00b2 increases from 0.009 at the 1-day horizon to 0.014 at 5 days but then declines to 0.011 at 10 days. The initial increase likely reflects signal accumulation over the first week, while the subsequent decline suggests signal decay as other information sources\u2014earnings releases, macroeconomic data, and news flow\u2014dominate returns at longer horizons. The 5-day horizon thus appears to represent the optimal window for skew-change information incorporation.")

add_subsection_heading("Quintile Analysis")

add_body("Figure 2 and Table 4 present the quintile results. Mean forward returns decrease monotonically from Q1 (skew flattening) to Q5 (skew steepening), producing a Q5\u2013Q1 spread of \u221255.55 bp at the 5-day horizon (t = \u22124.12, p < 0.001), representing an annualized return differential of approximately 28.9%.", indent=False)

add_figure_placeholder("[FIGURE 2: Bar chart of mean 5-day forward S&P 500 returns by \u0394Skew quintile with 95% CI error bars]")
add_figure_caption("Figure 2: Mean 5-day forward S&P 500 returns by \u0394Skew quintile. Q1 represents the strongest skew-flattening days (+24.13 bp), declining monotonically to Q5, the strongest skew-steepening days (\u221231.42 bp). Error bars represent 95% confidence intervals.")

add_table_cjsj(
    ["Quintile", "\u0394Skew Range", "1-Day (bp)", "5-Day (bp)", "10-Day (bp)"],
    [
        ["Q1 (Flat)", "< \u22120.48", "+5.82", "+24.13", "+41.27"],
        ["Q2", "\u22120.48 to \u22120.14", "+3.14", "+12.87", "+22.54"],
        ["Q3 (Neutral)", "\u22120.14 to +0.11", "+0.87", "+4.21", "+9.82"],
        ["Q4", "+0.11 to +0.51", "\u22121.43", "\u22127.34", "\u22128.14"],
        ["Q5 (Steep)", "> +0.51", "\u22126.92", "\u221231.42", "\u221242.83"],
        ["Q5\u2212Q1", "", "\u221212.74", "\u221255.55", "\u221284.10"],
    ],
    "Table 4: Mean forward returns by \u0394Skew quintile. The Q5\u2013Q1 spread is significant at p < 0.001 for all horizons."
)

add_subsection_heading("Volatility Regime Analysis")

add_body("Table 5 reports the 5-day regression results by VIX regime. The predictive power of \u0394Skew exhibits pronounced regime dependence. In low-VIX environments, the relationship is not statistically significant (p = 0.219). In high-VIX environments, the coefficient magnitude increases more than six-fold (\u03b2 = \u221238.47 vs. \u22126.14) with R\u00b2 = 4.1% (Figure 3).", indent=False)

add_table_cjsj(
    ["VIX Regime", "\u03b2 (bp)", "t-stat", "p-value", "R\u00b2", "N"],
    [
        ["Low (< 15)", "\u22126.14", "\u22121.23", "0.219", "0.006", "1,284"],
        ["Medium (15\u201325)", "\u221217.82", "\u22122.31", "0.021", "0.018", "2,134"],
        ["High (> 25)", "\u221238.47", "\u22123.14", "0.002", "0.041", "852"],
    ],
    "Table 5: Predictive regressions by volatility regime (5-day horizon). The steepening slope from low to high VIX illustrates regime dependence."
)

add_figure_placeholder("[FIGURE 3: Scatter plots of \u0394Skew vs. 5-day forward return by VIX regime with OLS regression lines]")
add_figure_caption("Figure 3: Scatter plots of \u0394Skew versus 5-day forward returns segmented by VIX regime. Panel A: Low VIX (<15). Panel B: Medium VIX (15\u201325). Panel C: High VIX (>25). OLS regression lines with 95% confidence bands are overlaid. The steepening slope from Panel A to Panel C illustrates the regime dependence of the predictive relationship.")

add_figure_placeholder("[FIGURE 4: Time series of daily \u0394Skew and S&P 500 index price (2008\u20132024)]")
add_figure_caption("Figure 4: Time series of daily \u0394Skew (top panel) and S&P 500 closing price (bottom panel) from January 2008 to December 2024. Shaded gray regions indicate high-VIX periods (VIX > 25). Notable \u0394Skew spikes are labeled with corresponding market events.")

add_subsection_heading("Directional Prediction and Robustness")

add_body("The logistic classifier achieved 55.7% directional accuracy at the 5-day horizon (AUC = 0.561), statistically significant but modest\u2014consistent with the general finding that even significant financial predictors explain a small fraction of return variation (14). The \u0394Skew signal proved robust across all specification variations: alternative delta levels (10-delta puts), smoothing windows (3-day rolling), exclusion of OpEx weeks, sub-period splits, and VIX term structure controls (Table 6).", indent=False)

add_table_cjsj(
    ["Specification", "\u03b2 (bp)", "t-stat", "p-value"],
    [
        ["Baseline (25\u0394 put)", "\u221218.73", "\u22122.89", "0.004"],
        ["10-delta put skew", "\u221222.41", "\u22122.47", "0.014"],
        ["3-day rolling \u0394Skew", "\u221214.82", "\u22122.68", "0.007"],
        ["Excl. OpEx weeks", "\u221219.14", "\u22122.77", "0.006"],
        ["2008\u20132015 only", "\u221221.37", "\u22122.14", "0.033"],
        ["2016\u20132024 only", "\u221215.92", "\u22122.03", "0.043"],
        ["Control VIX\u2013VIX3M", "\u221217.28", "\u22122.54", "0.011"],
    ],
    "Table 6: Robustness tests for the 5-day horizon regression coefficient. The negative predictive relationship persists across all specifications."
)

add_figure_placeholder("[FIGURE 5: Rolling 2-year \u03b2 coefficient with 95% confidence bands (2010\u20132024)]")
add_figure_caption("Figure 5: Rolling 2-year regression coefficient (\u03b2) of \u0394Skew on 5-day forward returns, estimated with an expanding window starting in 2010. The shaded region represents the 95% confidence band. The coefficient remains persistently negative across the full sample period.")

add_figure_placeholder("[FIGURE 6: Cumulative return of quintile-based strategy vs. buy-and-hold S&P 500]")
add_figure_caption("Figure 6: Cumulative return comparison: a quintile-based strategy (long S&P 500 on Q1 skew-flattening days, flat on Q5 skew-steepening days) versus buy-and-hold S&P 500, January 2008\u2013December 2024. Strategy outperformance is concentrated during high-volatility episodes.")

# ══════════════════════════════════════════════════════════════
# IV. DISCUSSION
# ══════════════════════════════════════════════════════════════
add_section_heading("IV. Discussion")

add_body("Our findings indicate that daily changes in S&P 500 options IV skew contain statistically significant and economically meaningful predictive information for short-term index returns. The negative relationship between skew steepening and forward returns is consistent across multiple frameworks and robust to specification changes.", indent=False)

add_body("These results support the informed trading hypothesis (15). When sophisticated market participants\u2014institutional hedgers, proprietary desks, and volatility arbitrageurs\u2014anticipate negative price movements, they preferentially express views through OTM put options, steepening skew before the underlying moves (6). Because options markets process information more rapidly due to leverage and lower capital requirements, this repositioning manifests as a lead signal.")

add_body("The pronounced regime dependence provides additional support. During low-volatility environments, options flows are dominated by routine hedging and yield-enhancement strategies, generating noise. During high-VIX periods, the signal-to-noise ratio improves because: (a) informed participants trade more aggressively under asymmetric risk (10), (b) dealer hedging flows amplify directional skew impacts (16), and (c) market maker inventory constraints become binding (17).")

add_body("The predictive power can also be understood through dealer gamma exposure. When end-users purchase OTM puts (steepening skew), dealers acquire negative gamma, requiring them to sell the underlying as it declines\u2014amplifying downward pressure (16). This mechanical hedging creates a self-reinforcing cycle, providing a non-informational channel through which skew changes predict returns. Both the informed trading and dealer mechanics channels likely contribute to the observed pattern.")

add_body("Our findings extend prior work on skew-level predictability (6) to skew changes at the index level, and align with Bollerslev, Tauchen, and Zhou (2009), who found regime-dependent predictive power in the variance risk premium (18). It is worth noting the relationship between \u0394Skew and the CBOE SKEW index: while the CBOE SKEW captures risk-neutral skewness from the full OTM option strip, \u0394Skew focuses on daily changes in the 25-delta/ATM spread\u2014a narrower but more responsive measure. Future work could examine whether combining these signals improves predictive power.")

add_body("The modest directional accuracy (53.8\u201355.7%) is consistent with financial econometrics: even significant predictors explain small fractions of return variation (14). Nevertheless, the 55.55 bp weekly quintile spread represents a meaningful signal for multi-factor quantitative strategies.")

add_subsection_heading("Limitations and Future Directions")

add_body("Several limitations should be acknowledged. First, R\u00b2 values (0.9\u20134.1%) indicate \u0394Skew explains only a small fraction of return variation. Second, this study uses end-of-day data; intraday skew dynamics may contain additional predictive content. Third, transaction costs and execution slippage are not modeled. Fourth, the sample covers a single asset (S&P 500); generalizability to other indices remains untested. Fifth, the 30-day constant-maturity interpolation relies on two bracketing expiration cycles; while both are observable at measurement time, the shifting interpolation weights introduce measurement noise that may attenuate the true predictive signal. Sixth, the CBOE SKEW index was not directly tested as an alternative predictor, limiting comparison against an established tail-risk benchmark.", indent=False)

add_body("Future research should extend the analysis to other major equity indices (NASDAQ-100, Russell 2000, Euro Stoxx 50), incorporate intraday options data to test time-of-day concentration effects, compare the delta-based \u0394Skew signal against daily changes in the CBOE SKEW index, combine \u0394Skew with complementary options signals (put-call ratios, implied correlation, VIX term structure) in multivariate frameworks, and apply machine learning methods (random forests, gradient-boosted trees) to capture nonlinear regime interactions.")

# ══════════════════════════════════════════════════════════════
# V. CONCLUSION
# ══════════════════════════════════════════════════════════════
add_section_heading("V. Conclusion")

add_body("This study provides evidence that daily changes in S&P 500 options implied volatility skew are statistically significant predictors of short-term index returns over the period 2008\u20132024. Skew steepening predicts negative forward returns, while skew flattening predicts positive returns. The relationship is monotonic across quintiles, robust to alternative specifications, and exhibits pronounced regime dependence with predictive power concentrated during high-volatility environments. These findings contribute to the literature on information transmission between options and equity markets and highlight the potential of options-derived signals as components of quantitative forecasting frameworks.", indent=False)

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
add_section_heading("References")

refs = [
    '[1]  E. F. Fama, "Efficient Capital Markets: A Review of Theory and Empirical Work," J. Finance, vol. 25, no. 2, pp. 383\u2013417, 1970, doi: 10.2307/2325486.',
    '[2]  S. Chakravarty, H. Gulen, and S. Mayhew, "Informed Trading in Stock and Option Markets," J. Finance, vol. 59, no. 3, pp. 1235\u20131258, 2004, doi: 10.1111/j.1540-6261.2004.00661.x.',
    '[3]  B. Dumas, J. Fleming, and R. E. Whaley, "Implied Volatility Functions: Empirical Tests," J. Finance, vol. 53, no. 6, pp. 2059\u20132106, 1998, doi: 10.1111/0022-1082.00083.',
    '[4]  M. Rubinstein, "Implied Binomial Trees," J. Finance, vol. 49, no. 3, pp. 771\u2013818, 1994, doi: 10.1111/j.1540-6261.1994.tb00079.x.',
    '[5]  N. G\u00e2rleanu, L. H. Pedersen, and A. M. Poteshman, "Demand-Based Option Pricing," Rev. Financ. Stud., vol. 22, no. 10, pp. 4259\u20134299, 2009, doi: 10.1093/rfs/hhp005.',
    '[6]  Y. Xing, X. Zhang, and R. Zhao, "What Does the Individual Option Volatility Smirk Tell Us About Future Equity Returns?" J. Financ. Quant. Anal., vol. 45, no. 3, pp. 641\u2013662, 2010, doi: 10.1017/S0022109010000220.',
    '[7]  M. Cremers and D. Weinbaum, "Deviations from Put-Call Parity and Stock Return Predictability," J. Financ. Quant. Anal., vol. 45, no. 2, pp. 335\u2013367, 2010, doi: 10.1017/S002210901000013X.',
    '[8]  T. G. Bali and A. Hovakimian, "Volatility Spreads and Expected Stock Returns," Manage. Sci., vol. 55, no. 11, pp. 1797\u20131812, 2009, doi: 10.1287/mnsc.1090.1063.',
    '[9]  P. Dennis and S. Mayhew, "Risk-Neutral Skewness: Evidence from Stock Options," J. Financ. Quant. Anal., vol. 37, no. 3, pp. 471\u2013493, 2002, doi: 10.2307/3594989.',
    '[10] A. Buraschi and J. Jackwerth, "The Price of a Smile: Hedging and Spanning in Option Markets," Rev. Financ. Stud., vol. 14, no. 2, pp. 495\u2013527, 2001, doi: 10.1093/rfs/14.2.495.',
    '[11] C. B. Mixon, "The Implied Volatility Term Structure of Stock Index Options," J. Empir. Finance, vol. 14, no. 3, pp. 333\u2013354, 2007, doi: 10.1016/j.jempfin.2006.06.003.',
    '[12] R. E. Whaley, "Understanding the VIX," J. Portfolio Manage., vol. 35, no. 3, pp. 98\u2013105, 2009, doi: 10.3905/JPM.2009.35.3.098.',
    '[13] W. K. Newey and K. D. West, "A Simple, Positive Semi-Definite, Heteroskedasticity and Autocorrelation Consistent Covariance Matrix," Econometrica, vol. 55, no. 3, pp. 703\u2013708, 1987, doi: 10.2307/1913610.',
    '[14] J. Y. Campbell and S. B. Thompson, "Predicting Excess Stock Returns Out of Sample: Can Anything Beat the Historical Average?" Rev. Financ. Stud., vol. 21, no. 4, pp. 1509\u20131531, 2008, doi: 10.1093/rfs/hhm055.',
    '[15] K. Back, "Asymmetric Information and Options," Rev. Financ. Stud., vol. 6, no. 3, pp. 435\u2013472, 1993, doi: 10.1093/rfs/6.3.435.',
    '[16] M. K. Brunnermeier and L. H. Pedersen, "Market Liquidity and Funding Liquidity," Rev. Financ. Stud., vol. 22, no. 6, pp. 2201\u20132238, 2009, doi: 10.1093/rfs/hhn098.',
    '[17] S. X. Ni, N. D. Pearson, and A. M. Poteshman, "Stock Price Clustering on Option Expiration Dates," J. Financ. Econ., vol. 78, no. 1, pp. 49\u201387, 2005, doi: 10.1016/j.jfineco.2004.08.005.',
    '[18] T. Bollerslev, G. Tauchen, and H. Zhou, "Expected Stock Returns and Variance Risk Premia," Rev. Financ. Stud., vol. 22, no. 11, pp. 4463\u20134492, 2009, doi: 10.1093/rfs/hhp008.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Split bracket number from rest
    bracket_end = ref.index(']') + 1
    run = p.add_run(ref[:bracket_end])
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True

    rest = ref[bracket_end:]
    # Italicize journal abbreviation (text between first and second comma after quotes)
    run = p.add_run(rest)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)

# === SAVE ===
out = '/Users/siddarthakodithyala/NHSJS_Manuscript_CJSJ_Format.docx'
doc.save(out)
print(f"Done! Saved to {out}")
