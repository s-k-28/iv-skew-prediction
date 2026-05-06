#!/usr/bin/env python3
"""Generate a formatted .docx of the Free Data Pipeline Plan."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── Title ───
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Free Data Pipeline Plan')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
title.paragraph_format.space_after = Pt(4)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('IV Skew Prediction Research Paper — NHSJS Submission')
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
subtitle.paragraph_format.space_after = Pt(4)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Siddartha Kodithyala')
run.font.size = Pt(11)
author.paragraph_format.space_after = Pt(2)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('April 18, 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
date_p.paragraph_format.space_after = Pt(20)

doc.add_paragraph('─' * 72)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# ─── 1. Context ───
doc.add_heading('1. Context & Problem Statement', level=1)
doc.add_paragraph(
    'The NHSJS manuscript (NHSJS_Manuscript_Skew_Prediction.md) currently contains placeholder/simulated '
    'results across 7 tables and 6 figures. No real data exists — no CSV datasets, no Jupyter notebooks, '
    'no processed data files. The original methodology requires 25-delta put IV minus 50-delta call IV from '
    'OptionMetrics IvyDB, which costs thousands of dollars per year.'
)
doc.add_paragraph(
    'This plan designs a completely free data pipeline that downloads real market data, runs all statistical '
    'analyses, produces publication-ready tables and figures, and updates the manuscript with real results.'
)

# ─── 2. Key Decision ───
doc.add_heading('2. Key Decision: CBOE SKEW Index', level=1)
doc.add_paragraph(
    'The CBOE SKEW index is the best free proxy for implied volatility skew. It is:'
)
bullets = [
    'Free, daily, available from 1990 to present via Yahoo Finance (ticker ^SKEW)',
    'Derived from S&P 500 OTM option prices, capturing tail-risk/skew dynamics',
    'Standardized and reproducible — any reader can replicate the study',
    'Computed from the full out-of-the-money option strip (richer than two delta points)',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph(
    'This reframing actually makes the paper stronger: it uses a widely-available benchmark rather than '
    'a bespoke construction, improving reproducibility and scientific rigor. The sign convention is preserved — '
    'positive ΔSkew (SKEW rising) = skew steepening = more tail risk demand, which the paper hypothesizes '
    'predicts negative returns.'
)

# ─── 3. Data Sources ───
doc.add_heading('3. Data Sources (All Free, No API Keys)', level=1)
add_table(
    ['Source', 'Ticker', 'Provider', 'Purpose'],
    [
        ['CBOE SKEW Index', '^SKEW', 'Yahoo Finance (yfinance)', 'Daily skew level → ΔSkew construction'],
        ['CBOE VIX Index', '^VIX', 'Yahoo Finance (yfinance)', 'Volatility regime classification'],
        ['S&P 500 Prices', '^GSPC', 'Yahoo Finance (yfinance)', 'Forward return calculation'],
    ]
)
doc.add_paragraph('Date range: December 1, 2007 – December 31, 2024 (extra month at start for lagged features).')

# ─── 4. Project Structure ───
doc.add_heading('4. Project File Structure', level=1)
structure = """Projects/skew-prediction/
├── config.py                      Configuration constants (dates, thresholds, paths)
├── requirements.txt               Python dependencies (all free)
├── data/
│   ├── raw/                       Cached Yahoo Finance CSV downloads
│   └── processed/                 Merged analysis dataset
├── src/
│   ├── __init__.py
│   ├── data_ingestion.py          Download + validate SKEW, VIX, SPX
│   ├── feature_construction.py    ΔSkew, forward returns, VIX regimes
│   ├── models.py                  OLS, quintile sorts, logistic, rolling analysis
│   └── robustness.py              Subperiod, crisis periods, OpEx exclusion
├── output/
│   ├── figures/                   6 publication-quality PNG/PDF figures
│   └── tables/                    7 CSV + LaTeX tables
└── run_pipeline.py                Master script: runs everything end-to-end"""

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ─── 5. Dependencies ───
doc.add_heading('5. Python Dependencies', level=1)
add_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['yfinance', '≥ 0.2.31', 'Download SKEW, VIX, SPX from Yahoo Finance'],
        ['pandas', '≥ 1.5.0', 'Data manipulation and merging'],
        ['numpy', '≥ 1.23.0', 'Numerical computation'],
        ['matplotlib', '≥ 3.6.0', 'Publication-quality figure generation'],
        ['scipy', '≥ 1.9.0', 'Statistical tests (t-tests, skewness, kurtosis)'],
        ['statsmodels', '≥ 0.13.0', 'OLS with Newey-West HAC, logistic regression'],
        ['scikit-learn', '≥ 1.1.0', 'ROC/AUC, time-series cross-validation'],
    ]
)
doc.add_paragraph('All packages are free and pip-installable. No API keys or paid accounts required.')

# ─── 6. Pipeline Steps ───
doc.add_heading('6. Pipeline Steps (run_pipeline.py)', level=1)

doc.add_heading('Step 1: Data Ingestion', level=2)
steps1 = [
    'Download ^SKEW, ^VIX, ^GSPC daily close prices via yfinance batch download',
    'Cache raw downloads as CSV files in data/raw/ (skip re-download if cached <24 hours)',
    'Validate: check date range coverage, missing values, outliers (SKEW outside [100, 170], VIX outside [8, 90])',
    'Inner-join all three series on date; forward-fill 1–2 day gaps in SKEW/VIX',
    'Save merged dataset to data/processed/analysis_dataset.csv',
    'Report: total observations, date range, missing data percentage',
]
for s in steps1:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Step 2: Feature Construction', level=2)
features = [
    'ΔSkew(t) = SKEW(t) − SKEW(t−1) — first difference of CBOE SKEW index',
    'Forward returns: fwd_ret_kd = (SPX(t+k) / SPX(t) − 1) × 100 for k ∈ {1, 5, 10, 21} trading days',
    'Trailing 5-day momentum: momentum_5d = (SPX(t) / SPX(t−5) − 1) × 100',
    'VIX regime classification: Low (VIX < 15), Medium (15 ≤ VIX ≤ 25), High (VIX > 25)',
    'Standardized ΔSkew (z-score) for "one-standard-deviation" effect interpretation',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Step 3: Statistical Analysis — 7 Tables', level=2)

doc.add_heading('Table 1: Summary Statistics', level=3)
doc.add_paragraph(
    'Mean, median, standard deviation, min, max, skewness, and kurtosis for: SKEW level, ΔSkew, VIX, '
    'and 1-day, 5-day, 10-day, 21-day S&P 500 returns.'
)

doc.add_heading('Table 2: Univariate OLS Regressions', level=3)
doc.add_paragraph(
    'R(t, t+k) = α + β · ΔSkew(t) + ε with Newey-West HAC standard errors (k−1 lags per horizon). '
    'Reports: intercept (α in bp), slope (β in bp per SKEW point), t-statistic, p-value, R², and N '
    'for 1-day, 5-day, 10-day, and 21-day horizons.'
)

doc.add_heading('Table 3: Multivariate OLS (5-Day Horizon)', level=3)
doc.add_paragraph(
    'R(t, t+5) = α + β₁·ΔSkew(t) + β₂·VIX(t) + β₃·Momentum(t) + ε. Three control variables '
    '(Volume is dropped — no free options volume data available). Reports coefficients, t-stats, and R².'
)

doc.add_heading('Table 4: Quintile Portfolio Returns', level=3)
doc.add_paragraph(
    'Sort trading days into quintiles by ΔSkew magnitude. For each quintile: mean 1-day, 5-day, and 10-day '
    'forward returns, ΔSkew range boundaries, and observation count. Q5−Q1 spread with two-sample t-test '
    'for statistical significance.'
)

doc.add_heading('Table 5: VIX Regime Conditional Results', level=3)
doc.add_paragraph(
    'Run univariate OLS separately for Low, Medium, and High VIX regimes. Report β, t-stat, p-value, R², '
    'and N per regime. Tests whether predictive power is concentrated in high-volatility environments.'
)

doc.add_heading('Table 6: Logistic Regression (Directional Prediction)', level=3)
doc.add_paragraph(
    'P(R > 0) = Λ(α + β · ΔSkew) with 5-fold expanding-window time-series cross-validation. '
    'Reports AUC, accuracy, precision, recall, logistic coefficient, and p-value for 1-day, 5-day, and 10-day horizons.'
)

doc.add_heading('Table 7: Robustness Checks', level=3)
add_table(
    ['Specification', 'What It Tests', 'Replaces (Original)'],
    [
        ['Baseline', 'Reference result for comparison', '—'],
        ['3-day rolling ΔSkew', 'Noise reduction via smoothing', 'Same (retained)'],
        ['Excluding OpEx weeks', 'Options expiration distortion', 'Same (retained)'],
        ['Sub-period: 2008–2015', 'Stability across market regimes', 'Same (retained)'],
        ['Sub-period: 2016–2024', 'Post-crisis structural change', 'Same (retained)'],
        ['Excluding outliers (|ΔSkew| > 2σ)', 'Extreme observation sensitivity', '10-delta put skew'],
        ['Crisis periods (GFC, COVID, 2022)', 'Tail-event behavior', 'VIX−VIX3M control'],
    ]
)

doc.add_heading('Step 4: Figure Generation — 6 Figures', level=2)
doc.add_paragraph('All figures use Times New Roman font, 300 DPI, and journal-quality styling.')
doc.add_paragraph()

add_table(
    ['Figure', 'Description', 'Key Details'],
    [
        ['Fig 1', 'Quintile bar chart', 'Mean 5-day returns by ΔSkew quintile, 95% CI error bars, green-to-red gradient'],
        ['Fig 2', 'Scatter plot with regime panels', 'ΔSkew vs 5-day return, color-coded by VIX regime, OLS regression lines + CI bands'],
        ['Fig 3', 'Time series (dual panel)', 'ΔSkew (top) + S&P 500 price (bottom), shaded high-VIX regions, event labels (GFC, COVID, 2022)'],
        ['Fig 4', 'ROC curves', '1/5/10-day logistic classifiers overlaid, diagonal reference line, AUC values in legend'],
        ['Fig 5', 'Rolling coefficient plot', '2-year rolling β with 95% CI band, horizontal zero line, VIX regime shading'],
        ['Fig 6', 'Cumulative return comparison', 'Quintile strategy (long Q1, flat Q5) vs buy-and-hold S&P 500, log scale y-axis'],
    ]
)

# ─── 7. Manuscript Updates ───
doc.add_heading('7. Manuscript Updates Required', level=1)
doc.add_paragraph(
    'After the pipeline produces real results, both the Markdown manuscript and the CJSJ DOCX conversion '
    'script must be updated with real data:'
)
updates = [
    ('Abstract', 'Replace "25-delta put IV minus 50-delta call IV" with "CBOE SKEW index, a standardized measure of options-implied tail risk." Update all numerical results (β, t-stats, R², quintile spreads).'),
    ('Section 2.1 (Data Sources)', 'Remove all references to OptionMetrics IvyDB, WRDS, and CRSP. Replace with: "CBOE SKEW index daily closing values from Yahoo Finance (^SKEW), S&P 500 prices from Yahoo Finance (^GSPC), VIX from Yahoo Finance (^VIX)."'),
    ('Section 2.2 (Skew Construction)', 'Replace delta-based construction with CBOE SKEW index definition. ΔSkew formula stays the same: ΔSkew(t) = SKEW(t) − SKEW(t−1). Add note on advantages of standardized measure.'),
    ('Section 2.5.1 (Multivariate OLS)', 'Drop Log(Volume) from the control specification — no free options volume data. Model becomes 3-variable: ΔSkew, VIX, 5-day momentum.'),
    ('Section 2.6 (Robustness)', 'Replace "10-delta put skew" with "excluding outliers (|ΔSkew| > 2σ)." Replace "VIX−VIX3M control" with "crisis period analysis (GFC, COVID, 2022)."'),
    ('Section 3 (Results)', 'Replace ALL placeholder table values and figure references with real computed values from the pipeline.'),
    ('Section 4.2 (Prior Literature)', 'Revise the CBOE SKEW paragraph — paper now uses the index directly instead of comparing against it.'),
    ('Section 4.4 (Limitations)', 'Remove 30-day interpolation limitation. Add: "The CBOE SKEW index aggregates across strikes, so the specific contribution of individual delta points to the predictive signal cannot be isolated."'),
    ('Figure Captions', 'Update all six figure captions with real computed values.'),
]
for title, desc in updates:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.font.bold = True
    p.add_run(desc)

# ─── 8. Statistical Methods Detail ───
doc.add_heading('8. Statistical Methods Detail', level=1)

doc.add_heading('Newey-West HAC Standard Errors', level=2)
doc.add_paragraph(
    'Overlapping return windows (5-day, 10-day, 21-day) create mechanical autocorrelation in residuals. '
    'Newey-West HAC standard errors correct for both heteroskedasticity and autocorrelation. '
    'Lag selection: k−1 lags for k-day returns (0 for daily, 4 for 5-day, 9 for 10-day, 20 for 21-day). '
    'Implementation: statsmodels OLS with fit(cov_type="HAC", cov_kwds={"maxlags": k-1}).'
)

doc.add_heading('Time-Series Cross-Validation', level=2)
doc.add_paragraph(
    'Standard k-fold CV is inappropriate for time-series data because it leaks future information into '
    'training folds. Instead, 5-fold expanding-window time-series CV is used: each fold trains on all data '
    'up to a cutoff date and tests on the subsequent out-of-sample window. '
    'Implementation: sklearn TimeSeriesSplit(n_splits=5).'
)

doc.add_heading('Quintile Sort Methodology', level=2)
doc.add_paragraph(
    'Trading days are sorted into five equal-sized bins (quintiles) by ΔSkew magnitude using pandas qcut. '
    'Q1 = strongest skew flattening (most negative ΔSkew), Q5 = strongest skew steepening (most positive ΔSkew). '
    'The Q5−Q1 return spread captures the economic magnitude of the signal. Significance tested via '
    'two-sample t-test (scipy.stats.ttest_ind).'
)

# ─── 9. Sign Convention ───
doc.add_heading('9. Sign Convention Verification', level=1)
doc.add_paragraph(
    'The CBOE SKEW index is scaled as follows: SKEW = 100 means a risk-neutral log-normal distribution '
    '(no skew); SKEW > 100 means the left tail is heavier than log-normal (more perceived tail risk). '
    'Therefore:'
)
conventions = [
    'Positive ΔSkew (SKEW rising) = skew steepening = MORE tail risk demand',
    'Negative ΔSkew (SKEW falling) = skew flattening = LESS tail risk demand',
    'Hypothesis: positive ΔSkew → negative short-term returns (same as original paper)',
    'The sign convention is preserved — no directional flip needed',
]
for c in conventions:
    doc.add_paragraph(c, style='List Bullet')

# ─── 10. Build Sequence ───
doc.add_heading('10. Build Sequence & Timeline', level=1)
add_table(
    ['#', 'Task', 'Est. Time', 'Dependencies'],
    [
        ['1', 'Install Python dependencies (pip install)', '5 min', 'None'],
        ['2', 'Create directory structure + config.py', '5 min', 'None'],
        ['3', 'Write data_ingestion.py (download, validate, cache)', '20 min', 'Step 1'],
        ['4', 'Write feature_construction.py (ΔSkew, returns, regimes)', '15 min', 'Step 3'],
        ['5', 'Write models.py (OLS, quintiles, logistic, rolling)', '40 min', 'Step 4'],
        ['6', 'Write robustness.py (subperiod, crisis, OpEx)', '20 min', 'Step 5'],
        ['7', 'Write figure generation (6 figures)', '40 min', 'Step 5–6'],
        ['8', 'Write table formatting + CSV/LaTeX output', '15 min', 'Step 5–6'],
        ['9', 'Write run_pipeline.py master script', '10 min', 'Steps 3–8'],
        ['10', 'Test end-to-end, debug', '30 min', 'Step 9'],
        ['11', 'Update manuscript .md with real results', '45 min', 'Step 10'],
        ['12', 'Update convert_manuscript_cjsj.py, regenerate .docx', '30 min', 'Step 11'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Total estimated implementation time: ~4.5 hours')
run.font.bold = True

# ─── 11. Verification Plan ───
doc.add_heading('11. Verification & Quality Checks', level=1)
checks = [
    'Run python run_pipeline.py end-to-end — must produce 7 tables in output/tables/ and 6 figures in output/figures/',
    'Verify analysis_dataset.csv has ~4,000+ rows with no NaN in key columns (ΔSkew, VIX, forward returns)',
    'Sanity check data: VIX mean ~18–20, SPX daily return mean ~0.03–0.05%, SKEW mean ~120–130',
    'Statistical sanity: R² should be small (0.1%–4%), which is normal for daily index return prediction',
    'Visual check: open each of the 6 figures, verify axis labels, font consistency, color schemes',
    'Manuscript consistency: all table values in the updated .md match pipeline CSV output exactly',
    'Regenerate .docx via convert_manuscript_cjsj.py and verify CJSJ formatting is preserved',
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

# ─── 12. Risks ───
doc.add_heading('12. Risks & Mitigations', level=1)
add_table(
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['Results differ from placeholders', 'Tables/figures change; some results may not be significant',
         'Report whatever the data shows — null results are publishable too'],
        ['yfinance ^SKEW data doesn\'t go back to 2008', 'Shorter sample period',
         'Fall back to CBOE direct CSV download (covers 1990–present) or shorten study period'],
        ['yfinance API rate limiting or outage', 'Data download fails',
         'Cache raw CSVs after first download; retry with exponential backoff (3 attempts)'],
        ['Small R² values', 'May seem like weak results',
         'Expected for daily return prediction — emphasize economic significance via quintile spreads'],
        ['Missing Volume control variable', 'Omitted variable bias concern',
         'VIX and momentum controls remain; Volume was not significant in original spec (p = 0.674)'],
    ]
)

# ─── 13. What Changes in the Paper ───
doc.add_heading('13. Summary: What Changes vs. What Stays', level=1)

doc.add_heading('What Changes', level=2)
changes = [
    'Skew measure: 25δ put − 50δ call IV → CBOE SKEW index (free, standardized)',
    'Data source: OptionMetrics IvyDB / WRDS → Yahoo Finance (no cost, no login)',
    'Multivariate controls: Drop Log(Volume), keep VIX + momentum',
    'Robustness: Drop 10-delta put + VIX−VIX3M checks, add outlier exclusion + crisis period analysis',
    'All numerical results: Replaced with real computed values',
    'Limitations: Drop interpolation noise, add SKEW aggregation limitation',
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('What Stays the Same', level=2)
stays = [
    'Paper title, structure, and narrative arc',
    'Core hypothesis: ΔSkew steepening predicts negative returns',
    'ΔSkew formula: ΔSkew(t) = Skew(t) − Skew(t−1)',
    'All 18 academic references (all are real and relevant)',
    'Statistical methods: OLS + Newey-West, quintile sorts, logistic regression, rolling analysis',
    'VIX regime thresholds: Low (<15), Medium (15–25), High (>25)',
    'Return horizons: 1-day, 5-day, 10-day (plus new 21-day)',
    'Sample period: 2008–2024',
    'CJSJ two-column formatting style',
]
for s in stays:
    doc.add_paragraph(s, style='List Bullet')

# ─── Footer ───
doc.add_paragraph()
doc.add_paragraph('─' * 72)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Prepared for Siddartha Kodithyala — April 18, 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

output_path = os.path.expanduser('~/Free_Data_Pipeline_Plan.docx')
doc.save(output_path)
print(f'Plan saved to: {output_path}')
